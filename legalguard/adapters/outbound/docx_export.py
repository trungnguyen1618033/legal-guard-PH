"""Xuất Bản ghi nhớ sửa đổi (Phase C) ra Word .docx — adapter opt-in (cần `python-docx`).

Tách khỏi domain (domain chỉ tạo memo thuần). `python-docx` ở group `export` → import LƯỜI; chưa cài →
`DocxUnavailable` để caller trả lỗi rõ (markdown vẫn dùng được). `memo_to_docx(memo_dict) -> bytes`.
"""
from __future__ import annotations

import io


class DocxUnavailable(RuntimeError):
    """python-docx chưa được cài (group `export`)."""


def memo_to_docx(memo: dict) -> bytes:
    """memo (dict từ `AnalysisService.compile_memo`) → bytes .docx. Bảng: Điều|Vấn đề|Tính chất|Căn cứ|
    Đề xuất|Ưu tiên. Raise DocxUnavailable nếu thiếu lib."""
    try:
        from docx import Document
    except ImportError as exc:  # noqa: TRY003
        raise DocxUnavailable("Cần cài: uv sync --group export (python-docx).") from exc

    _STATUS = {"illegal": "TRÁI LUẬT (có thể vô hiệu)", "unfavorable": "Bất lợi"}
    _PRIO = {"must_fix": "Phải sửa", "negotiate": "Thương lượng", "acceptable": "Chấp nhận được"}
    doc = Document()
    doc.add_heading(memo.get("title") or "BẢN GHI NHỚ SỬA ĐỔI HỢP ĐỒNG", level=0)
    rows = memo.get("rows", [])
    cols = ["#", "Điều khoản", "Vấn đề", "Tính chất", "Căn cứ pháp lý", "Đề xuất sửa", "Ưu tiên"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for c, name in zip(table.rows[0].cells, cols):
        c.text = name
    for i, r in enumerate(rows, 1):
        nature = _STATUS.get(r.get("legal_status"), r.get("legal_status", ""))
        if r.get("legal_status") == "illegal" and r.get("violated_law"):
            nature += f" — {r['violated_law']}"
        cells = table.add_row().cells
        vals = [str(i), r.get("clause", ""), r.get("issue", ""), nature,
                r.get("legal_basis", ""), r.get("suggestion", ""),
                _PRIO.get(r.get("priority"), r.get("priority", ""))]
        for cell, v in zip(cells, vals):
            cell.text = v or "—"
    illegal = memo.get("illegal_count", sum(r.get("legal_status") == "illegal" for r in rows))
    doc.add_paragraph(f"Tổng: {len(rows)} điều khoản — {illegal} TRÁI LUẬT, {len(rows) - illegal} bất lợi.")
    doc.add_paragraph("⚠️ Bản ghi nhớ do AI hỗ trợ soạn — luật sư cần đối chiếu bản gốc trước khi sử dụng.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def redline_to_docx(rl: dict) -> bytes:
    """Bản ĐỐI CHIẾU sửa đổi (dict từ `compile_redline`) → .docx: mỗi điều khoản → tiêu đề + điều khoản CŨ
    (đỏ, gạch ngang) → điều khoản MỚI (xanh, highlight) song ngữ + căn cứ. Raise DocxUnavailable nếu thiếu lib."""
    try:
        from docx import Document
        from docx.enum.text import WD_COLOR_INDEX
        from docx.shared import RGBColor
    except ImportError as exc:  # noqa: TRY003
        raise DocxUnavailable("Cần cài: uv sync --group export (python-docx).") from exc

    _RED, _GREEN = RGBColor(0xC0, 0x00, 0x00), RGBColor(0x0A, 0x6A, 0x30)
    doc = Document()
    doc.add_heading(rl.get("title") or "BẢN ĐỐI CHIẾU SỬA ĐỔI HỢP ĐỒNG", level=0)
    if rl.get("protected_party"):
        doc.add_paragraph(f"Bên được bảo vệ: {rl['protected_party']}")

    def _run(par, text, *, color=None, strike=False, highlight=False, bold=False):
        r = par.add_run(text)
        if color is not None:
            r.font.color.rgb = color
        r.font.strike = strike
        r.bold = bold
        if highlight:
            r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        return r

    for i, row in enumerate(rl.get("rows", []), 1):
        tag = ""
        if row.get("legal_status") == "illegal":
            tag = f" — TRÁI LUẬT{(' (' + row['violated_law'] + ')') if row.get('violated_law') else ''}"
        doc.add_heading(f"({i}) {row.get('clause', '')}{tag}", level=2)
        if row.get("old"):
            p = doc.add_paragraph()
            _run(p, "Điều khoản gốc: ", bold=True)
            _run(p, row["old"], color=_RED, strike=True)      # cũ: đỏ + gạch ngang
        new_vi, new_en = row.get("new_vi", ""), row.get("new_en", "")
        if new_vi or new_en:
            p = doc.add_paragraph()
            _run(p, "Đề xuất sửa: ", bold=True)
            if new_vi:
                _run(p, new_vi, color=_GREEN, highlight=True)  # mới: xanh + highlight
            if new_en:
                pe = doc.add_paragraph()
                _run(pe, "EN: ", bold=True)
                _run(pe, new_en, color=_GREEN)
        if row.get("reason"):
            p = doc.add_paragraph()
            _run(p, "Căn cứ: ", bold=True)
            _run(p, row["reason"])
    doc.add_paragraph()
    doc.add_paragraph("⚠️ Bản đối chiếu do AI hỗ trợ soạn — luật sư cần đối chiếu bản gốc trước khi áp dụng.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_STATUS_TAG = {"illegal": "⚠️ TRÁI LUẬT", "unfavorable": "Bất lợi"}


def comment_to_docx(doc_data: dict) -> bytes:
    """File Word có COMMENT thật (bong bóng nhận xét Word) — mỗi điều khoản bị gắn cờ → đoạn trích nguyên
    văn + 1 comment chứa: [trạng thái] rủi ro (+ điều luật vi phạm) · đề xuất sửa song ngữ · căn cứ.
    Giống 'file có comment như ChatGPT' nhưng nhận xét GROUNDED (từ dữ liệu case).

    `doc_data` = {title, protected_party, contract_type, items:[{clause, evidence, risk, legal_status,
    violated_law, vi, en, rationale}]}. Raise DocxUnavailable nếu thiếu python-docx hoặc bản cũ không có
    add_comment (cần python-docx ≥ 1.2)."""
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError as exc:  # noqa: TRY003
        raise DocxUnavailable("Cần cài: uv sync --group export (python-docx).") from exc

    doc = Document()
    if not hasattr(doc, "add_comment"):     # python-docx < 1.2 chưa hỗ trợ comment
        raise DocxUnavailable("Cần python-docx ≥ 1.2 để chèn comment vào Word.")

    _RED = RGBColor(0xC0, 0x00, 0x00)
    doc.add_heading(doc_data.get("title") or "HỢP ĐỒNG — BẢN RÀ SOÁT CÓ NHẬN XÉT", level=0)
    if doc_data.get("protected_party"):
        doc.add_paragraph(f"Bên được bảo vệ: {doc_data['protected_party']}")
    if doc_data.get("contract_type"):
        doc.add_paragraph(f"Loại hợp đồng: {doc_data['contract_type']}")
    doc.add_paragraph("Mỗi điều khoản dưới đây có gắn NHẬN XÉT (comment) — mở bằng Microsoft Word để xem "
                      "chi tiết rủi ro và đề xuất sửa.")
    doc.add_paragraph()

    items = sorted(doc_data.get("items") or [],
                   key=lambda it: it.get("legal_status") != "illegal")   # TRÁI LUẬT lên đầu
    for i, it in enumerate(items, 1):
        clause = (it.get("clause") or "").strip()
        evidence = (it.get("evidence") or "").strip() or clause
        status = it.get("legal_status", "unfavorable")
        doc.add_heading(f"({i}) {clause}", level=2)
        p = doc.add_paragraph()
        run = p.add_run(evidence)                      # đoạn trích nguyên văn — mỏ neo cho comment
        if status == "illegal":
            run.font.color.rgb = _RED
        # Nội dung comment: trạng thái + rủi ro + điều luật vi phạm + đề xuất song ngữ + căn cứ.
        parts = [f"[{_STATUS_TAG.get(status, 'Bất lợi')}]"]
        if (risk := (it.get("risk") or "").strip()):
            parts.append(risk.rstrip(".") + ".")
        if status == "illegal" and (vl := (it.get("violated_law") or "").strip()):
            parts.append(f"Trái quy định tại {vl} — phần vi phạm có thể bị tuyên vô hiệu.")
        if (vi := (it.get("vi") or "").strip()):
            parts.append(f"Đề xuất sửa (VI): {vi}")
        if (en := (it.get("en") or "").strip()):
            parts.append(f"Suggested (EN): {en}")
        if (rat := (it.get("rationale") or "").strip()):
            parts.append(f"Căn cứ: {rat.rstrip('.')}.")
        doc.add_comment(runs=p.runs, text="\n".join(parts), author="Legal Guard", initials="LG")

    doc.add_paragraph()
    doc.add_paragraph("⚠️ Nhận xét do AI hỗ trợ soạn — luật sư cần đối chiếu bản gốc trước khi áp dụng.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
