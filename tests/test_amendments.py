"""Phase C — gộp Bản ghi nhớ sửa đổi (memo) + xuất Word."""
from legalguard.domain.amendments import compile_memo, compile_redline

_ITEMS = [
    {"clause": "Điều 5 — Trọng tài", "risk": "Bất lợi địa điểm", "legal_status": "unfavorable",
     "legal_basis": "", "suggestion": "Đổi sang VIAC", "priority": "negotiate"},
    {"clause": "Điều 2 — Phạt 15%", "risk": "Vượt trần", "legal_status": "illegal",
     "violated_law": "Điều 301 LTM 2005", "suggestion": "Về 8%", "priority": "must_fix"},
]


def test_compile_memo_sorts_illegal_first_and_builds_table():
    memo = compile_memo(_ITEMS, title="Memo HĐ X", protected_party="Bên Mua")
    # illegal lên đầu (đòn mạnh nhất)
    assert memo.rows[0].legal_status == "illegal" and memo.rows[0].clause.startswith("Điều 2")
    assert memo.illegal_count == 1
    md = memo.markdown
    assert "# Memo HĐ X" in md and "Bên Mua" in md
    assert "| Điều khoản | Vấn đề |" in md          # có bảng
    assert "⚖️ TRÁI LUẬT (Điều 301 LTM 2005)" in md  # nhãn + điều luật vi phạm
    assert "luật sư cần đối chiếu" in md             # disclaimer


def test_compile_memo_skips_empty_clause_and_escapes_pipe():
    memo = compile_memo([{"clause": "", "risk": "x"},
                         {"clause": "Điều 3 | a", "risk": "ghi chú | pipe"}])
    assert len(memo.rows) == 1                       # bỏ mục thiếu clause
    assert "Điều 3 / a" in memo.markdown             # '|' trong ô → '/' (không vỡ bảng Markdown)


def test_memo_to_docx_bytes():
    # python-docx (group export) → .docx hợp lệ (zip 'PK'). Bỏ qua nếu chưa cài.
    from dataclasses import asdict

    import pytest
    pytest.importorskip("docx")
    from legalguard.adapters.outbound.docx_export import memo_to_docx
    data = memo_to_docx(asdict(compile_memo(_ITEMS)))
    assert data[:2] == b"PK" and len(data) > 500     # docx = zip


# ---- REDLINE (Mức 1 sửa-file): bản đối chiếu cũ→mới ----
_RL_ITEMS = [
    {"clause": "Điều 8 — Thanh toán", "evidence": "Thanh toán trong 90 ngày.",
     "vi": "Thanh toán trong 30 ngày.", "en": "Payment within 30 days.", "legal_status": "unfavorable"},
    {"clause": "Điều 5 — Phạt", "evidence": "Phạt 15% giá trị hợp đồng.", "vi": "Phạt tối đa 8%.",
     "en": "Cap 8%.", "rationale": "Điều 301 LTM 2005", "legal_status": "illegal",
     "violated_law": "Điều 301"},
]


def test_compile_redline_sorts_illegal_first_flexible_keys():
    rl = compile_redline(_RL_ITEMS, title="Đối chiếu HĐ", protected_party="Bên B")
    assert rl.rows[0].legal_status == "illegal" and rl.rows[0].clause.startswith("Điều 5")  # illegal đầu
    assert rl.rows[0].old == "Phạt 15% giá trị hợp đồng." and rl.rows[0].new_vi == "Phạt tối đa 8%."
    assert rl.rows[0].reason == "Điều 301 LTM 2005" and rl.illegal_count == 1
    assert rl.rows[1].new_en == "Payment within 30 days."      # 'en' → new_en
    assert compile_redline([{"clause": ""}]).rows == []        # bỏ mục thiếu clause


def test_redline_to_docx_bytes_and_content():
    from dataclasses import asdict

    import pytest
    pytest.importorskip("docx")
    import io

    from docx import Document

    from legalguard.adapters.outbound.docx_export import redline_to_docx
    data = redline_to_docx(asdict(compile_redline(_RL_ITEMS)))
    assert data[:2] == b"PK" and len(data) > 500              # docx = zip
    doc = Document(io.BytesIO(data))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Phạt 15% giá trị hợp đồng." in full and "Phạt tối đa 8%." in full   # cũ + mới
    # điều khoản CŨ phải gạch ngang (strike) — dấu hiệu redline
    assert any(r.font.strike for p in doc.paragraphs for r in p.runs)
